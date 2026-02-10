from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Predefined data variables
applicant_data = {
    'title': 'Mr',
    'first_name': 'RAHUL',
    'middle_name': 'KUMAR',
    'last_name': 'SHARMA',
    'college_address': 'National Dairy Research Institute, Karnal, Haryana',
    'college_pincode': '132001',
    'residential_address': 'House No. 45, Sector 12, Model Town, Karnal, Haryana',
    'residential_pincode': '132001',
    'mailing_address_type': 'College',  # or 'Residence'
    'phone': '0184-2252800',
    'mobile': '9876543210',
    'email': 'rahul.sharma@example.com',
    'course': 'B.Tech Dairy Technology',
    'year': '3rd Year',
    'dob_day': '15',
    'dob_month': 'March',
    'dob_year': '2005',
    'course_name': 'Bachelor of Technology in Dairy Technology',
    'duration_from_month': 'August',
    'duration_from_year': '2022',
    'duration_to_month': 'May',
    'duration_to_year': '2026',
    'utr_no': 'DD123456789',
    'payment_date': '05/02/2026',
    'bank_name': 'State Bank of India'
}

# Qualifications data (as list of dictionaries for table)
qualifications = [
    {
        'degree': '12th (Senior Secondary)',
        'university': 'CBSE Board',
        'year': '2022'
    },
    {
        'degree': '10th (Secondary)',
        'university': 'CBSE Board',
        'year': '2020'
    }
]

def fill_template(template_path, output_path):
    """
    Fill the DOCX template with predefined data
    """
    # Load the template
    doc = Document(template_path)
    
    # Replace placeholders in the document
    replace_placeholders(doc, applicant_data)
    
    # Find and fill tables
    fill_qualifications_table(doc, qualifications)
    
    # Save the filled document
    doc.save(output_path)
    print(f"Document successfully created: {output_path}")

def replace_placeholders(doc, data):
    """
    Replace {{placeholder}} text in the entire document
    """
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = '{{' + key + '}}'
            if placeholder in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if placeholder in inline[i].text:
                        inline[i].text = inline[i].text.replace(placeholder, str(value))
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder = '{{' + key + '}}'
                    if placeholder in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(value))

def fill_qualifications_table(doc, qualifications_data):
    """
    Fill the qualifications table with data
    """
    # Find the table with {{qualifications_table}} marker
    for table in doc.tables:
        if any('{{qualifications_table}}' in cell.text for row in table.rows for cell in row.cells):
            # Clear the marker row if it exists
            for i, row in enumerate(table.rows):
                if any('{{qualifications_table}}' in cell.text for cell in row.cells):
                    # This is the data insertion point
                    # Remove marker text
                    for cell in row.cells:
                        if '{{qualifications_table}}' in cell.text:
                            cell.text = ''
                    
                    # Add qualification rows
                    for idx, qual in enumerate(qualifications_data):
                        if idx == 0:
                            # Use the existing row
                            current_row = row
                        else:
                            # Add new rows
                            current_row = table.add_row()
                        
                        # Fill the cells
                        current_row.cells[0].text = qual['degree']
                        current_row.cells[1].text = qual['university']
                        current_row.cells[2].text = qual['year']
                    break

def create_form_from_scratch(output_path):
    """
    Alternative: Create the entire form from scratch without template
    """
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('INDIAN DAIRY ASSOCIATION\nSTUDENT MEMBERSHIP APPLICATION FORM')
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    # 1. Name of Applicant
    table1 = doc.add_table(rows=4, cols=2)
    table1.style = 'Table Grid'
    
    table1.rows[0].cells[0].text = '1. Name of Applicant\n   (in capital letters)'
    table1.rows[0].cells[1].text = f"Title: {applicant_data['title']}"
    
    table1.rows[1].cells[0].text = ''
    table1.rows[1].cells[1].text = f"First Name: {applicant_data['first_name']}"
    
    table1.rows[2].cells[0].text = ''
    table1.rows[2].cells[1].text = f"Middle Name: {applicant_data['middle_name']}"
    
    table1.rows[3].cells[0].text = ''
    table1.rows[3].cells[1].text = f"Last Name: {applicant_data['last_name']}"
    
    doc.add_paragraph()
    
    # 2. College Address
    table2 = doc.add_table(rows=2, cols=2)
    table2.style = 'Table Grid'
    table2.rows[0].cells[0].text = '2. College Address'
    table2.rows[0].cells[1].text = applicant_data['college_address']
    table2.rows[1].cells[0].text = ''
    table2.rows[1].cells[1].text = f"Pin Code: {applicant_data['college_pincode']}"
    
    doc.add_paragraph()
    
    # 3. Residential Address
    table3 = doc.add_table(rows=2, cols=2)
    table3.style = 'Table Grid'
    table3.rows[0].cells[0].text = 'Residential Address'
    table3.rows[0].cells[1].text = applicant_data['residential_address']
    table3.rows[1].cells[0].text = ''
    table3.rows[1].cells[1].text = f"Pin Code: {applicant_data['residential_pincode']}"
    
    doc.add_paragraph()
    
    # 4. Contact Details
    table4 = doc.add_table(rows=3, cols=2)
    table4.style = 'Table Grid'
    table4.rows[0].cells[0].text = 'Mailing Address:\nContact Details'
    table4.rows[0].cells[1].text = f"Mailing Address Type: {applicant_data['mailing_address_type']}"
    table4.rows[1].cells[0].text = ''
    table4.rows[1].cells[1].text = f"Phone: {applicant_data['phone']}    Mobile: {applicant_data['mobile']}"
    table4.rows[2].cells[0].text = ''
    table4.rows[2].cells[1].text = f"E-mail: {applicant_data['email']}"
    
    doc.add_paragraph()
    
    # 5. Qualifications
    p = doc.add_paragraph()
    p.add_run('3. Qualifications').bold = True
    
    table5 = doc.add_table(rows=len(qualifications) + 1, cols=3)
    table5.style = 'Table Grid'
    
    # Header row
    table5.rows[0].cells[0].text = 'Degree/Diploma'
    table5.rows[0].cells[1].text = 'University/Institution'
    table5.rows[0].cells[2].text = 'Years Awarded'
    
    # Make header bold
    for cell in table5.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Data rows
    for idx, qual in enumerate(qualifications, start=1):
        table5.rows[idx].cells[0].text = qual['degree']
        table5.rows[idx].cells[1].text = qual['university']
        table5.rows[idx].cells[2].text = qual['year']
    
    doc.add_paragraph()
    
    # 6. Present Status
    table6 = doc.add_table(rows=1, cols=2)
    table6.style = 'Table Grid'
    table6.rows[0].cells[0].text = f"4. Present Status - Course: {applicant_data['course']}"
    table6.rows[0].cells[1].text = f"Year: {applicant_data['year']}"
    
    doc.add_paragraph()
    
    # 7. Date of Birth
    table7 = doc.add_table(rows=2, cols=3)
    table7.style = 'Table Grid'
    table7.rows[0].cells[0].text = '5. Date of Birth'
    table7.rows[0].cells[1].text = ''
    table7.rows[0].cells[2].text = ''
    table7.rows[1].cells[0].text = f"Day: {applicant_data['dob_day']}"
    table7.rows[1].cells[1].text = f"Month: {applicant_data['dob_month']}"
    table7.rows[1].cells[2].text = f"Year: {applicant_data['dob_year']}"
    
    doc.add_paragraph()
    
    # 8. Course Name and Duration
    p = doc.add_paragraph()
    p.add_run(f"6. Name of the Course: {applicant_data['course_name']}")
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f"   Duration of Course:\n   From: {applicant_data['duration_from_month']} {applicant_data['duration_from_year']}   To: {applicant_data['duration_to_month']} {applicant_data['duration_to_year']}")
    
    doc.add_paragraph()
    
    # Declaration
    declaration = doc.add_paragraph()
    declaration.add_run('Declaration: ').bold = True
    declaration.add_run('The above information is true to the best of my knowledge and belief. If admitted to the Association I undertake to abide by the Constitution of the Association as contained therein or as amended.')
    
    doc.add_paragraph()
    doc.add_paragraph(f"Date: {applicant_data['payment_date']}")
    doc.add_paragraph().add_run('(Signature of the applicant)').italic = True
    
    doc.add_paragraph()
    
    # Eligibility Criteria
    p = doc.add_paragraph()
    p.add_run('7. Eligibility Criteria:').bold = True
    
    p1 = doc.add_paragraph('• Any person who is a student for diploma or bachelor\'s degree in any stream of dairy science and not employed by any organisation or drawing any salary, is eligible for under-graduate student membership.', style='List Bullet')
    
    p2 = doc.add_paragraph('• Any person who is a student for post-graduate or Doctorate Programme in any stream of dairy science and not employed by any organisation or drawing any salary is eligible for student membership.', style='List Bullet')
    
    doc.add_paragraph().add_run('P.T.O.').italic = True
    
    # Payment Information
    doc.add_page_break()
    
    p = doc.add_paragraph()
    p.add_run('8. Payment Details').bold = True
    
    doc.add_paragraph(f"UTR No. / DD No. / Cheque at par: {applicant_data['utr_no']}")
    doc.add_paragraph(f"Date: {applicant_data['payment_date']}")
    doc.add_paragraph(f"Name of the Bank: {applicant_data['bank_name']}")
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run('BANK DETAILS: ').bold = True
    p.add_run('Name: Indian Dairy Association; SB a/c No: 90562170000024; IFSC: CNRB0019009;\nBank: Canara Bank; Branch Address: Delhi Tamil Sangam Building, Sector-V, R.K. Puram, New Delhi.')
    
    doc.add_paragraph()
    
    table8 = doc.add_table(rows=1, cols=2)
    table8.style = 'Table Grid'
    table8.rows[0].cells[0].text = 'STUDENT MEMBERSHIP FEE (PER COURSE) @ (Rs. 700 + GST@18%)'
    table8.rows[0].cells[1].text = 'TOTAL: Rs. 826/- Including GST'
    
    # Save document
    doc.save(output_path)
    print(f"Form created successfully: {output_path}")

if __name__ == "__main__":
    # Method 1: Using a template (recommended)
    # Make sure you have ST_Nation.docx template in the same directory
    template_file = "COPY.docx"
    output_file = "Filled_Membership_Form.docx"
    
    try:
        fill_template(template_file, output_file)
    except FileNotFoundError:
        print(f"Template file '{template_file}' not found.")
        print("Creating form from scratch instead...")
        create_form_from_scratch(output_file)
    
    # Method 2: Create from scratch (alternative if no template)
    # create_form_from_scratch("Membership_Form_Generated.docx")
