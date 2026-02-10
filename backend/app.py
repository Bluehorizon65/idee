from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
import os
import json
import tempfile
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import time

app = Flask(__name__)

# Production CORS configuration
CORS(app, resources={
    r"/api/*": {
        "origins": [
            "https://stellular-daifuku-971e29.netlify.app",  # Netlify production domain
            "http://localhost:3000",  # Development
            "http://localhost:5173",   # Vite dev
            "http://localhost:5000"    # Local Flask dev
        ],
        "methods": ["GET", "POST", "OPTIONS"],
        "allow_headers": ["Content-Type"],
        "max_age": 3600
    }
})

def create_membership_form_from_json(data, output_path, logo_path='IDA_logo.png'):
    """
    Create membership form from JSON data received from frontend
    """
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # =====================================================================
    # HEADER: Logo at top center, then Title with Photo Box
    # =====================================================================
    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    try:
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(logo_path, width=Inches(3), height=Inches(2))
    except:
        run = logo_paragraph.add_run('[IDA LOGO - 3"x2"]')
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        run.italic = True
    
    # Title and Photo box table
    title_photo_table = doc.add_table(rows=1, cols=3)
    title_photo_table.autofit = False
    title_photo_table.allow_autofit = False
    
    title_photo_table.columns[0].width = Inches(0.5)
    title_photo_table.columns[1].width = Inches(4.5)
    title_photo_table.columns[2].width = Inches(2.0)
    
    # Remove borders
    for row in title_photo_table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for edge in ('top', 'left', 'bottom', 'right'):
                edge_element = OxmlElement(f'w:{edge}')
                edge_element.set(qn('w:val'), 'none')
                tcBorders.append(edge_element)
            tcPr.append(tcBorders)
    
    # Title cell
    title_cell = title_photo_table.rows[0].cells[1]
    title_paragraph = title_cell.paragraphs[0]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = title_paragraph.add_run('APPLICATION FORM FOR STUDENT MEMBERSHIP\n')
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    
    run = title_paragraph.add_run('(FOR INDIAN CITIZEN)')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    # Photo box
    photo_cell = title_photo_table.rows[0].cells[2]
    photo_paragraph = photo_cell.paragraphs[0]
    photo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    tc = photo_cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    left_node = OxmlElement('w:left')
    left_node.set(qn('w:w'), '720')
    left_node.set(qn('w:type'), 'dxa')
    tcMar.append(left_node)
    for margin_name in ['top', 'bottom', 'right']:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), '0')
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)
    
    photo_box_table = photo_cell.add_table(rows=1, cols=1)
    photo_box_table.style = 'Table Grid'
    photo_box_cell = photo_box_table.rows[0].cells[0]
    photo_box_cell.width = Inches(1.5)
    
    p = photo_box_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n\n\nPlease\naffix your\nphotograph\n\n\n')
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run.italic = True
    
    doc.add_paragraph()
    
    # =====================================================================
    # TABLE 1: NAME OF APPLICANT
    # =====================================================================
    table1 = doc.add_table(rows=4, cols=2)
    table1.style = 'Table Grid'
    table1.columns[0].width = Inches(2.5)
    table1.columns[1].width = Inches(4.5)
    
    cell = table1.rows[0].cells[0]
    p = cell.paragraphs[0]
    run = p.add_run('1. Name of Applicant\n   (in capital letters)')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    cell = table1.rows[0].cells[1]
    p = cell.paragraphs[0]
    run = p.add_run(f"Title: {data.get('title', '')}")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    table1.rows[1].cells[1].text = f"First Name: {data.get('first_name', '')}"
    table1.rows[2].cells[1].text = f"Middle Name: {data.get('middle_name', '')}"
    table1.rows[3].cells[1].text = f"Last Name: {data.get('last_name', '')}"
    
    for row in table1.rows[1:]:
        for paragraph in row.cells[1].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # =====================================================================
    # TABLE 2: COLLEGE ADDRESS
    # =====================================================================
    table2 = doc.add_table(rows=2, cols=2)
    table2.style = 'Table Grid'
    table2.columns[0].width = Inches(2.5)
    table2.columns[1].width = Inches(4.5)
    
    cell = table2.rows[0].cells[0]
    p = cell.paragraphs[0]
    run = p.add_run('2. College Address')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    table2.rows[0].cells[1].text = data.get('college_address', '')
    table2.rows[1].cells[1].text = f"Pin Code: {data.get('college_pincode', '')}"
    
    for row in table2.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # =====================================================================
    # TABLE 3: RESIDENTIAL ADDRESS
    # =====================================================================
    table3 = doc.add_table(rows=2, cols=2)
    table3.style = 'Table Grid'
    table3.columns[0].width = Inches(2.5)
    table3.columns[1].width = Inches(4.5)
    
    cell = table3.rows[0].cells[0]
    p = cell.paragraphs[0]
    run = p.add_run('Residential Address')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    table3.rows[0].cells[1].text = data.get('residential_address', '')
    table3.rows[1].cells[1].text = f"Pin Code: {data.get('residential_pincode', '')}"
    
    for row in table3.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # =====================================================================
    # MAILING ADDRESS & CONTACT DETAILS
    # =====================================================================
    p = doc.add_paragraph()
    run = p.add_run('Mailing Address:               College ')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    mailing_college = '☑' if data.get('mailing_address_type') == 'College' else '☐'
    mailing_residence = '☑' if data.get('mailing_address_type') == 'Residence' else '☐'
    
    run = p.add_run(mailing_college)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    run = p.add_run('          Residence ')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    run = p.add_run(mailing_residence)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    # Contact Details
    p = doc.add_paragraph()
    run = p.add_run('Contact Details             Phone')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    run = p.add_run(' ' + data.get('phone', '') + '                 ')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.underline = True
    
    run = p.add_run('  Mobile')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    run = p.add_run(' ' + data.get('mobile', '') + '                 ')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.underline = True
    
    # Email
    p = doc.add_paragraph()
    run = p.add_run('                                     E-mail')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    run = p.add_run(' ' + data.get('email', '') + '                                         ')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.underline = True
    
    # =====================================================================
    # TABLE 5: QUALIFICATIONS
    # =====================================================================
    p = doc.add_paragraph()
    run = p.add_run('3. Qualifications')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    qualifications = data.get('qualifications', [])
    table5 = doc.add_table(rows=len(qualifications) + 1, cols=3)
    table5.style = 'Table Grid'
    
    table5.columns[0].width = Inches(2.3)
    table5.columns[1].width = Inches(2.7)
    table5.columns[2].width = Inches(2.0)
    
    headers = ['Degree/Diploma', 'University/Institution', 'Years Awarded']
    for idx, header in enumerate(headers):
        cell = table5.rows[0].cells[idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
    
    for idx, qual in enumerate(qualifications, start=1):
        table5.rows[idx].cells[0].text = qual.get('degree', '')
        table5.rows[idx].cells[1].text = qual.get('university', '')
        table5.rows[idx].cells[2].text = qual.get('year', '')
        
        for cell in table5.rows[idx].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # =====================================================================
    # TABLE 6: PRESENT STATUS
    # =====================================================================
    table6 = doc.add_table(rows=1, cols=2)
    table6.style = 'Table Grid'
    table6.columns[0].width = Inches(4.5)
    table6.columns[1].width = Inches(2.5)
    
    cell = table6.rows[0].cells[0]
    p = cell.paragraphs[0]
    run = p.add_run('4. Present Status')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    p.add_run(f"\n   Course: {data.get('course', '')}")
    
    cell = table6.rows[0].cells[1]
    p = cell.paragraphs[0]
    run = p.add_run(f"Year: {data.get('year', '')}")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    for cell in table6.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # =====================================================================
    # TABLE 7: DATE OF BIRTH
    # =====================================================================
    table7 = doc.add_table(rows=2, cols=4)
    table7.style = 'Table Grid'
    
    cell = table7.rows[0].cells[0]
    merged_cell = cell.merge(table7.rows[0].cells[3])
    p = merged_cell.paragraphs[0]
    run = p.add_run('5. Date of Birth')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    dob_headers = ['Day', 'Month', 'Year']
    dob_values = [data.get('dob_day', ''), data.get('dob_month', ''), data.get('dob_year', '')]
    
    for idx in range(3):
        cell = table7.rows[1].cells[idx]
        p = cell.paragraphs[0]
        run = p.add_run(f"{dob_headers[idx]}\n{dob_values[idx]}")
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # =====================================================================
    # COURSE INFORMATION
    # =====================================================================
    p = doc.add_paragraph()
    run = p.add_run('6. Name of the Course: ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    run = p.add_run(data.get('course_name', ''))
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    run = p.add_run('   (Please attach your identity card attested by Principal OR Dean of the College as proof of studentship.)')
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('   Duration of Course:')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    run = p.add_run(f"   From: {data.get('duration_from_month', '')} {data.get('duration_from_year', '')}   ")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    run = p.add_run('To: ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    run = p.add_run(f"{data.get('duration_to_month', '')} {data.get('duration_to_year', '')}")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # =====================================================================
    # DECLARATION
    # =====================================================================
    p = doc.add_paragraph()
    run = p.add_run('Declaration: ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    run = p.add_run('The above information is true to the best of my knowledge and belief. If admitted to the Association I undertake to abide by the Constitution of the Association as contained therein or as amended.')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run(f"Date: {data.get('payment_date', '')}")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('(Signature of the applicant)')
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # =====================================================================
    # ELIGIBILITY CRITERIA
    # =====================================================================
    p = doc.add_paragraph()
    run = p.add_run('7. Eligibility Criteria:')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run('Any person who is a student for diploma or bachelor\'s degree in any stream of dairy science and not employed by any organisation or drawing any salary, is eligible for under-graduate student membership.')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run('Any person who is a student for post-graduate or Doctorate Programme in any stream of dairy science and not employed by any organisation or drawing any salary is eligible for student membership.')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('P.T.O.')
    run.italic = True
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    # =====================================================================
    # PAGE BREAK
    # =====================================================================
    doc.add_page_break()
    
    # =====================================================================
    # PAGE 2: PAYMENT DETAILS
    # =====================================================================
    p = doc.add_paragraph()
    run = p.add_run('8. ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    run = p.add_run('The application form should be duly filled and returned to the concerned zone / IDA HQ directly along with the initial membership fee, paid by NEFT / Bank Draft / Cheque at par only. The Student Membership fee needs to be paid one-time only when joining course. It needs renewal for the next course.')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run(f"UTR No. / DD No. / Cheque at par: {data.get('utr_no', '')}       ")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    run = p.add_run(f"Date: {data.get('payment_date', '')}       ")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    run = p.add_run(f"Name of the Bank: {data.get('bank_name', '')}")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('BANK DETAILS: ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    run = p.add_run('Name: Indian Dairy Association; SB a/c No: 90562170000024; IFSC: CNRB0019009;')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    run = p.add_run('Bank: Canara Bank; Branch Address: Delhi Tamil Sangam Building, Sector-V, R.K. Puram, New Delhi.')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # =====================================================================
    # TABLE 8: FEE STRUCTURE
    # =====================================================================
    table8 = doc.add_table(rows=1, cols=2)
    table8.style = 'Table Grid'
    table8.columns[0].width = Inches(4.5)
    table8.columns[1].width = Inches(2.5)
    
    cell = table8.rows[0].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('STUDENT MEMBERSHIP FEE (PER COURSE)\n@ (Rs. 700 + GST@18%)')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    cell = table8.rows[0].cells[1]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TOTAL: Rs. 826/-\nIncluding GST')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # =====================================================================
    # JOURNAL NOTE
    # =====================================================================
    p = doc.add_paragraph()
    run = p.add_run('9. ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    run = p.add_run('Only Soft copy of the journal will be shared on the registered Email ID.')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    # UPI Payment Image
    try:
        upi_paragraph = doc.add_paragraph()
        upi_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        upi_run = upi_paragraph.add_run()
        upi_run.add_picture('payment.png', width=Inches(4.5))
    except:
        pass
    
    # =====================================================================
    # IDA OFFICE CONTACT INFORMATION
    # =====================================================================
    p = doc.add_paragraph()
    run = p.add_run('* IDA HQ: ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    run = p.add_run('IDA House, Sector-IV, R.K. Puram, New Delhi-110 022  Phones: 26182454, 26179781, 26165355')
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    run = p.add_run('  E-mail: admin@indairyasso.org / idahq@rediffmail.com  Web: www.indairyasso.org')
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    run = p.add_run('* South Zone: ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    run = p.add_run('The Secretary, (Address) Ph.: 080-25710661 Fax: 080-25710161')
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    run = p.add_run('* West Zone: ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    run = p.add_run('The Secretary, A-501, Dynasty Business Park, Andheri-Kurla Road, Andheri (East), Mumbai 400059')
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    run = p.add_run('  Email: chairman@idawz.org Website: www.idawz.org')
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    run = p.add_run('* North Zone: ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    run = p.add_run('The Secretary, IDA (NZ), IDA House, Sector IV, R.K. Puram, New Delhi - 110 022 Phones: 011-26170781, 26165355')
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    run = p.add_run('* East Zone: ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    run = p.add_run('The Secretary, c/o NDDB, Block-DK, Sector-II, Salt Lake, Kolkata-700 091 Phones: 033-23345877')
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    doc.save(output_path)

@app.route('/api/generate-form', methods=['POST'])
def generate_form():
    try:
        data = request.json
        
        # Create a temporary file
        timestamp = str(int(time.time() * 1000))
        filename = f'IDA_Membership_{timestamp}.docx'
        filepath = os.path.join(tempfile.gettempdir(), filename)
        
        # Generate the document
        create_membership_form_from_json(data, filepath)
        
        # Send the file
        return send_file(
            filepath,
            as_attachment=True,
            download_name=f'IDA_Student_Membership_{data.get("first_name", "Form")}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        # Delete the file after sending
        try:
            if os.path.exists(filepath):
                # Wait a bit before deleting to ensure file is sent
                time.sleep(1)
                os.remove(filepath)
        except:
            pass

@app.route('/api/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok'})

if __name__ == '__main__':
    print("=" * 70)
    print("IDA MEMBERSHIP FORM GENERATOR - FLASK SERVER")
    print("=" * 70)
    print("\nServer starting on http://localhost:5000")
    print("Frontend should connect to: http://localhost:5000/api/generate-form")
    print("\nPress CTRL+C to stop the server")
    print("=" * 70)
    # Production: use environment PORT, Development: use 5000
    port = int(os.environ.get('PORT', 5000))
    debug_mode = os.environ.get('FLASK_ENV') == 'development'
    app.run(debug=debug_mode, port=port, host='0.0.0.0')
