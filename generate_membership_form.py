from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Predefined data variables - CUSTOMIZE THESE
applicant_data = {
    'title': 'Mr',
    'first_name': 'RAHUL',
    'middle_name': 'KUMAR',
    'last_name': 'SHARMA',
    'college_address': 'National Dairy Research Institute, Karnal, Haryana',
    'college_pincode': '132001',
    'residential_address': 'House No. 45, Sector 12, Model Town, Karnal, Haryana',
    'residential_pincode': '132001',
    'mailing_address_type': 'College',  # 'College' or 'Residence'
    'phone': '0184-2252800',
    'mobile': '9876543210',
    'email': 'rahul.sharma@example.com',
    'course': 'B.Tech Dairy Technology',
    'year': '3rd Year',
    'dob_day': '15',
    'dob_month': 'March',
    'dob_year': '2005',
    'course_name': 'Bachelor of Technology in Dairy Technology',
    'duration_from_month': 'August',
    'duration_from_year': '2022',
    'duration_to_month': 'May',
    'duration_to_year': '2026',
    'utr_no': 'DD123456789',
    'payment_date': '05/02/2026',
    'bank_name': 'State Bank of India'
}

# Qualifications data - ADD MORE AS NEEDED
qualifications = [
    {
        'degree': '12th (Senior Secondary)',
        'university': 'CBSE Board',
        'year': '2022'
    },
    {
        'degree': '10th (Secondary)',
        'university': 'CBSE Board',
        'year': '2020'
    }
]

def set_cell_border(cell, **kwargs):
    """
    Set cell borders with custom properties.
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    # Create borders element
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), 'single')
            edge_element.set(qn('w:sz'), '4')
            edge_element.set(qn('w:space'), '0')
            edge_element.set(qn('w:color'), '000000')
            tcBorders.append(edge_element)
    
    tcPr.append(tcBorders)

def create_membership_form(output_path, logo_path='IDA_logo.png'):
    """
    Create a complete Indian Dairy Association Student Membership Application Form
    """
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # =====================================================================
    # HEADER: Logo at top center, then Title with Photo Box
    # =====================================================================
    # Row 1: Logo centered at top
    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    try:
        logo_run = logo_paragraph.add_run()
        # Logo size: 3 inches x 3 inches
        logo_run.add_picture(logo_path, width=Inches(3), height=Inches(2))
    except:
        # If logo file not found, add text placeholder
        run = logo_paragraph.add_run('[IDA LOGO - 3"x3"]')
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        run.italic = True
    
    # Row 2: Title and Photo box - using 3 columns for better control
    title_photo_table = doc.add_table(rows=1, cols=3)
    title_photo_table.autofit = False
    title_photo_table.allow_autofit = False
    
    # Set column widths
    title_photo_table.columns[0].width = Inches(0.5)  # Left spacer
    title_photo_table.columns[1].width = Inches(4.5)  # Title column (centered)
    title_photo_table.columns[2].width = Inches(2.0)  # Photo box column (right)
    
    # Remove borders from title/photo table
    for row in title_photo_table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for edge in ('top', 'left', 'bottom', 'right'):
                edge_element = OxmlElement(f'w:{edge}')
                edge_element.set(qn('w:val'), 'none')
                tcBorders.append(edge_element)
            tcPr.append(tcBorders)
    
    # Column 1: Empty spacer
    # (leave it empty)
    
    # Column 2: Title text (centered)
    title_cell = title_photo_table.rows[0].cells[1]
    title_paragraph = title_cell.paragraphs[0]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add main title
    run = title_paragraph.add_run('APPLICATION FORM FOR STUDENT MEMBERSHIP\n')
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    
    # Add subtitle
    run = title_paragraph.add_run('(FOR INDIAN CITIZEN)')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    # Column 3: Photo box (right aligned)
    photo_cell = title_photo_table.rows[0].cells[2]
    photo_paragraph = photo_cell.paragraphs[0]
    photo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add left padding to push photo box to the right
    tc = photo_cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    # Add left margin (1440 twips = 1 inch, 720 = 0.5 inch)
    left_node = OxmlElement('w:left')
    left_node.set(qn('w:w'), '720')  # 0.5 inch left padding
    left_node.set(qn('w:type'), 'dxa')
    tcMar.append(left_node)
    # Minimal margins for other sides
    for margin_name in ['top', 'bottom', 'right']:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), '0')
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)
    
    # Create a bordered box for photo
    photo_box_table = photo_cell.add_table(rows=1, cols=1)
    photo_box_table.style = 'Table Grid'
    photo_box_cell = photo_box_table.rows[0].cells[0]
    photo_box_cell.width = Inches(1.5)
    
    p = photo_box_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n\n\nPlease\naffix your\nphotograph\n\n\n')
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run.italic = True
    
    doc.add_paragraph()
    
    # =====================================================================
    # TABLE 1: NAME OF APPLICANT
    # =====================================================================
    table1 = doc.add_table(rows=4, cols=2)
    table1.style = 'Table Grid'
    
    # Set column widths
    table1.columns[0].width = Inches(2.5)
    table1.columns[1].width = Inches(4.5)
    
    # Row 1: Title
    cell = table1.rows[0].cells[0]
    p = cell.paragraphs[0]
    run = p.add_run('1. Name of Applicant\n   (in capital letters)')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    cell = table1.rows[0].cells[1]
    p = cell.paragraphs[0]
    run = p.add_run(f"Title: {applicant_data['title']}")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    # Row 2: First Name
    table1.rows[1].cells[1].text = f"First Name: {applicant_data['first_name']}"
    for paragraph in table1.rows[1].cells[1].paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
    
    # Row 3: Middle Name
    table1.rows[2].cells[1].text = f"Middle Name: {applicant_data['middle_name']}"
    for paragraph in table1.rows[2].cells[1].paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
    
    # Row 4: Last Name
    table1.rows[3].cells[1].text = f"Last Name: {applicant_data['last_name']}"
    for paragraph in table1.rows[3].cells[1].paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # =====================================================================
    # TABLE 2: COLLEGE ADDRESS
    # =====================================================================
    table2 = doc.add_table(rows=2, cols=2)
    table2.style = 'Table Grid'
    table2.columns[0].width = Inches(2.5)
    table2.columns[1].width = Inches(4.5)
    
    cell = table2.rows[0].cells[0]
    p = cell.paragraphs[0]
    run = p.add_run('2. College Address')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    table2.rows[0].cells[1].text = applicant_data['college_address']
    table2.rows[1].cells[1].text = f"Pin Code: {applicant_data['college_pincode']}"
    
    for row in table2.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # =====================================================================
    # TABLE 3: RESIDENTIAL ADDRESS
    # =====================================================================
    table3 = doc.add_table(rows=2, cols=2)
    table3.style = 'Table Grid'
    table3.columns[0].width = Inches(2.5)
    table3.columns[1].width = Inches(4.5)
    
    cell = table3.rows[0].cells[0]
    p = cell.paragraphs[0]
    run = p.add_run('Residential Address')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    table3.rows[0].cells[1].text = applicant_data['residential_address']
    table3.rows[1].cells[1].text = f"Pin Code: {applicant_data['residential_pincode']}"
    
    for row in table3.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # =====================================================================
    # MAILING ADDRESS & CONTACT DETAILS (Not a table - formatted text)
    # =====================================================================
    # Mailing Address line with checkboxes
    p = doc.add_paragraph()
    run = p.add_run('Mailing Address:               College ')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    # Checkbox for mailing type
    mailing_college = '☑' if applicant_data['mailing_address_type'] == 'College' else '☐'
    mailing_residence = '☑' if applicant_data['mailing_address_type'] == 'Residence' else '☐'
    
    run = p.add_run(mailing_college)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    run = p.add_run('          Residence ')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    run = p.add_run(mailing_residence)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    # Contact Details line
    p = doc.add_paragraph()
    run = p.add_run('Contact Details             Phone')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    run = p.add_run(' ' + applicant_data['phone'] + '                 ')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.underline = True
    
    run = p.add_run('  Mobile')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    run = p.add_run(' ' + applicant_data['mobile'] + '                 ')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.underline = True
    
    # Email line
    p = doc.add_paragraph()
    run = p.add_run('                                     E-mail')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    run = p.add_run(' ' + applicant_data['email'] + '                                         ')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.underline = True
    
    # =====================================================================
    # TABLE 5: QUALIFICATIONS
    # =====================================================================
    p = doc.add_paragraph()
    run = p.add_run('3. Qualifications')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    table5 = doc.add_table(rows=len(qualifications) + 1, cols=3)
    table5.style = 'Table Grid'
    
    # Set column widths
    table5.columns[0].width = Inches(2.3)
    table5.columns[1].width = Inches(2.7)
    table5.columns[2].width = Inches(2.0)
    
    # Header row
    headers = ['Degree/Diploma', 'University/Institution', 'Years Awarded']
    for idx, header in enumerate(headers):
        cell = table5.rows[0].cells[idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
    
    # Data rows
    for idx, qual in enumerate(qualifications, start=1):
        table5.rows[idx].cells[0].text = qual['degree']
        table5.rows[idx].cells[1].text = qual['university']
        table5.rows[idx].cells[2].text = qual['year']
        
        for cell in table5.rows[idx].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # =====================================================================
    # TABLE 6: PRESENT STATUS
    # =====================================================================
    table6 = doc.add_table(rows=1, cols=2)
    table6.style = 'Table Grid'
    table6.columns[0].width = Inches(4.5)
    table6.columns[1].width = Inches(2.5)
    
    cell = table6.rows[0].cells[0]
    p = cell.paragraphs[0]
    run = p.add_run('4. Present Status')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    p.add_run(f"\n   Course: {applicant_data['course']}")
    
    cell = table6.rows[0].cells[1]
    p = cell.paragraphs[0]
    run = p.add_run(f"Year: {applicant_data['year']}")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    for cell in table6.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # =====================================================================
    # TABLE 7: DATE OF BIRTH
    # =====================================================================
    table7 = doc.add_table(rows=2, cols=4)
    table7.style = 'Table Grid'
    
    # Merge first row cells for title
    cell = table7.rows[0].cells[0]
    merged_cell = cell.merge(table7.rows[0].cells[3])
    p = merged_cell.paragraphs[0]
    run = p.add_run('5. Date of Birth')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    # Second row
    dob_headers = ['Day', 'Month', 'Year']
    dob_values = [applicant_data['dob_day'], applicant_data['dob_month'], applicant_data['dob_year']]
    
    for idx in range(3):
        cell = table7.rows[1].cells[idx]
        p = cell.paragraphs[0]
        run = p.add_run(f"{dob_headers[idx]}\n{dob_values[idx]}")
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # =====================================================================
    # COURSE INFORMATION
    # =====================================================================
    p = doc.add_paragraph()
    run = p.add_run('6. Name of the Course: ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    run = p.add_run(applicant_data['course_name'])
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    run = p.add_run('   (Please attach your identity card attested by Principal OR Dean of the College as proof of studentship.)')
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('   Duration of Course:')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    run = p.add_run(f"   From: {applicant_data['duration_from_month']} {applicant_data['duration_from_year']}   ")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    run = p.add_run('To: ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    run = p.add_run(f"{applicant_data['duration_to_month']} {applicant_data['duration_to_year']}")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # =====================================================================
    # DECLARATION
    # =====================================================================
    p = doc.add_paragraph()
    run = p.add_run('Declaration: ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    run = p.add_run('The above information is true to the best of my knowledge and belief. If admitted to the Association I undertake to abide by the Constitution of the Association as contained therein or as amended.')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run(f"Date: {applicant_data['payment_date']}")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('(Signature of the applicant)')
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # =====================================================================
    # ELIGIBILITY CRITERIA
    # =====================================================================
    p = doc.add_paragraph()
    run = p.add_run('7. Eligibility Criteria:')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run('Any person who is a student for diploma or bachelor\'s degree in any stream of dairy science and not employed by any organisation or drawing any salary, is eligible for under-graduate student membership.')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run('Any person who is a student for post-graduate or Doctorate Programme in any stream of dairy science and not employed by any organisation or drawing any salary is eligible for student membership.')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('P.T.O.')
    run.italic = True
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    # =====================================================================
    # PAGE BREAK
    # =====================================================================
    doc.add_page_break()
    
    # =====================================================================
    # PAGE 2: PAYMENT DETAILS
    # =====================================================================
    p = doc.add_paragraph()
    run = p.add_run('8. ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    run = p.add_run('The application form should be duly filled and returned to the concerned zone / IDA HQ directly along with the initial membership fee, paid by NEFT / Bank Draft / Cheque at par only. The Student Membership fee needs to be paid one-time only when joining course. It needs renewal for the next course.')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run(f"UTR No. / DD No. / Cheque at par: {applicant_data['utr_no']}       ")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    run = p.add_run(f"Date: {applicant_data['payment_date']}       ")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    run = p.add_run(f"Name of the Bank: {applicant_data['bank_name']}")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('BANK DETAILS: ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    run = p.add_run('Name: Indian Dairy Association; SB a/c No: 90562170000024; IFSC: CNRB0019009;')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    run = p.add_run('Bank: Canara Bank; Branch Address: Delhi Tamil Sangam Building, Sector-V, R.K. Puram, New Delhi.')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # =====================================================================
    # TABLE 8: FEE STRUCTURE
    # =====================================================================
    table8 = doc.add_table(rows=1, cols=2)
    table8.style = 'Table Grid'
    table8.columns[0].width = Inches(4.5)
    table8.columns[1].width = Inches(2.5)
    
    cell = table8.rows[0].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('STUDENT MEMBERSHIP FEE (PER COURSE)\n@ (Rs. 700 + GST@18%)')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    cell = table8.rows[0].cells[1]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TOTAL: Rs. 826/-\nIncluding GST')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # =====================================================================
    # JOURNAL NOTE
    # =====================================================================
    p = doc.add_paragraph()
    run = p.add_run('9. ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    run = p.add_run('Only Soft copy of the journal will be shared on the registered Email ID.')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    #doc.add_paragraph()
    
    # UPI Payment Image
    try:
        upi_paragraph = doc.add_paragraph()
        upi_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        upi_run = upi_paragraph.add_run()
        upi_run.add_picture('payment.png', width=Inches(4.5))
    except:
        # If UPI image not found, add placeholder text
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('[UPI Payment QR Code Image]')
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        run.italic = True
    
    #doc.add_paragraph()
    
    # =====================================================================
    # IDA OFFICE CONTACT INFORMATION
    # =====================================================================
    p = doc.add_paragraph()
    run = p.add_run('* IDA HQ: ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    run = p.add_run('IDA House, Sector-IV, R.K. Puram, New Delhi-110 022  Phones: 26182454, 26179781, 26165355')
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    run = p.add_run('  E-mail: admin@indairyasso.org / idahq@rediffmail.com  Web: www.indairyasso.org')
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    run = p.add_run('* South Zone: ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    run = p.add_run('The Secretary, (Address) Ph.: 080-25710661 Fax: 080-25710161')
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    run = p.add_run('* West Zone: ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    run = p.add_run('The Secretary, A-501, Dynasty Business Park, Andheri-Kurla Road, Andheri (East), Mumbai 400059')
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    run = p.add_run('  Email: chairman@idawz.org Website: www.idawz.org')
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    run = p.add_run('* North Zone: ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    run = p.add_run('The Secretary, IDA (NZ), IDA House, Sector IV, R.K. Puram, New Delhi - 110 022 Phones: 011-26170781, 26165355')
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    run = p.add_run('* East Zone: ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    run = p.add_run('The Secretary, c/o NDDB, Block-DK, Sector-II, Salt Lake, Kolkata-700 091 Phones: 033-23345877')
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    
    # Save the document
    doc.save(output_path)
    print(f"✓ Membership form successfully created: {output_path}")
    print(f"✓ Applicant: {applicant_data['first_name']} {applicant_data['middle_name']} {applicant_data['last_name']}")
    print(f"✓ Course: {applicant_data['course']}")
    print(f"✓ Total qualifications added: {len(qualifications)}")

if __name__ == "__main__":
    output_file = "IDA_Student_Membership_Form.docx"
    logo_file = "IDA_logo.png"  # Place your logo image file in the same directory
    
    print("=" * 70)
    print("INDIAN DAIRY ASSOCIATION - STUDENT MEMBERSHIP FORM GENERATOR")
    print("=" * 70)
    print("\nGenerating form with predefined data...")
    print(f"Output file: {output_file}")
    print(f"Logo file: {logo_file} (optional - will show placeholder if not found)")
    print()
    
    create_membership_form(output_file, logo_file)
    
    print()
    print("=" * 70)
    print("DONE! Open the file to view the completed form.")
    print("=" * 70)
    print("\nTo customize data:")
    print("1. Edit the 'applicant_data' dictionary in this script")
    print("2. Modify the 'qualifications' list to add/remove qualifications")
    print("3. Place 'IDA_logo.png' in the same folder to add the logo image")
    print("4. Run the script again to generate a new form")
